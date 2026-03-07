import json
import io
import base64
import time
from http.server import BaseHTTPRequestHandler
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont


def set_font(paragraph, font_name="Times New Roman", size=12, bold=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        run.bold = bold


def add_bold_para(doc, text, font_name="Times New Roman", size=12, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = True
    set_font(p, font_name=font_name, size=size, bold=True)
    return p


def add_normal_para(doc, text, font_name="Times New Roman", size=12, align=None):
    if align is None:
        if any(c in text for c in ['*', '•', '·']) or '  ' in text:
            align = WD_ALIGN_PARAGRAPH.LEFT
        else:
            align = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(p, font_name=font_name, size=size)
    return p


def add_labeled_para(doc, label, content, font_name="Times New Roman", size=12):
    p = doc.add_paragraph()
    if any(c in content for c in ['*', '•', '·']) or '  ' in content:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p.add_run(f"{label} ")
    run_label.bold = True
    run_label.font.name = font_name
    run_label.font.size = Pt(size)
    run_content = p.add_run(content)
    run_content.font.name = font_name
    run_content.font.size = Pt(size)
    return p


def add_code_para(doc, code_text, font_name="Times New Roman", size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.name = font_name
    run.font.size = Pt(size)


def add_caption_para(doc, text, experiment_no, font_name="Times New Roman", size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {experiment_no} - {text}")
    run.font.name = font_name
    run.font.size = Pt(size)
    run.italic = False


def create_terminal_image(output_text, img_width=600):
    font_size = 14
    padding = 20
    
    # Cloud-safe fonts (Linux/Universal) + Windows Fallbacks
    font = None
    windir = os.environ.get('WINDIR', 'C:\\Windows')
    paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf",
        os.path.join(windir, 'Fonts', 'consola.ttf'),
        os.path.join(windir, 'Fonts', 'cour.ttf'),
        "DejaVuSansMono.ttf"
    ]
    
    for p in paths:
        try:
            if os.path.exists(p):
                font = ImageFont.truetype(p, font_size)
                break
        except: continue
        
    if font is None:
        font = ImageFont.load_default()

    lines = str(output_text).split('\n')
    line_height = font_size + 4 
    height = (len(lines) * line_height) + (2 * padding)
    
    img = Image.new('RGB', (img_width, height), color=(30, 30, 30))
    d = ImageDraw.Draw(img)
    
    y = padding
    for line in lines:
        try:
            d.text((padding, y), line.replace('\r', ''), font=font, fill=(210, 210, 210))
        except: pass
        y += line_height
        
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers.get("Content-Length", 0))
            body = self.rfile.read(content_length)
            data = json.loads(body)

            experiments = data.get("experiments", [])
            settings = data.get("settings", {})

            font_name = settings.get("fontName", "Times New Roman")
            body_size = int(settings.get("bodySize", 12))
            heading_size = int(settings.get("headingSize", 14))
            code_size = int(settings.get("codeSize", 10))
            caption_size = int(settings.get("captionSize", 10))
            image_width_inches = float(settings.get("imageWidth", 5.0))
            terminal_img_width = int(settings.get("terminalImgWidth", 600))
            output_filename = settings.get("outputFilename", "Generated_Practical_File.docx")

            doc = Document()

            for i, exp in enumerate(experiments, 1):
                aim = exp.get("aim", "")
                concept = exp.get("concept", "")
                code = exp.get("code", "")
                output = exp.get("output", "")
                caption = exp.get("caption", "")

                # Experiment heading
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Experiment No. {i}")
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(heading_size)
                doc.add_paragraph("")

                # Aim
                add_labeled_para(doc, "Aim:", aim, font_name, body_size)
                doc.add_paragraph("")

                # Concept
                add_labeled_para(doc, "Concept Used:", concept, font_name, body_size)
                doc.add_paragraph("")

                # Code
                add_bold_para(doc, "Code:", font_name, body_size)
                add_code_para(doc, code, font_name, code_size)
                doc.add_paragraph("")

                # Output
                add_bold_para(doc, "Output:", font_name, body_size)

                # Terminal image
                try:
                    img_buf = create_terminal_image(output, terminal_img_width)
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = pic_para.add_run()
                    run.add_picture(img_buf, width=Inches(image_width_inches))
                    add_caption_para(doc, caption, i, font_name, caption_size)
                except Exception as e:
                    add_normal_para(doc, f"[Error adding image: {e}]", font_name, body_size)
                    add_code_para(doc, output, font_name, code_size)

                if i < len(experiments):
                    doc.add_page_break()

            # Save to buffer
            file_buf = io.BytesIO()
            doc.save(file_buf)
            file_buf.seek(0)

            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f'attachment; filename="{output_filename}"')
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(file_buf.read())

        except Exception as e:
            self.send_response(500)
            self.send_header("Content-Type", "application/json")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(json.dumps({"error": str(e)}).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()
