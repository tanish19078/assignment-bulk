import json
import re
import io
import time
import os
import traceback
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from dotenv import load_dotenv

load_dotenv()
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


app = Flask(__name__, static_folder='public', static_url_path='')
CORS(app)

@app.route('/')
def serve_index():
    return send_from_directory('public', 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory('public', path)


# ==================== API: Parse Aims ====================
@app.route('/api/parse', methods=['POST'])
def api_parse():
    try:
        data = request.get_json()
        text = data.get('text', '')
        separator = data.get('separator', '---')

        pattern = r'\n\s*' + re.escape(separator) + r'+\s*\n'
        aim_blocks = re.split(pattern, text)
        aims = [b.strip() for b in aim_blocks if b.strip()]

        return jsonify({'aims': aims})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== STEP PARSING (OS MODE) ====================
def parse_steps(procedure_text):
    """Parse the procedure text into individual steps with explanation, command, and output.
    Handles both 'Step N:' format and 'N.' numbered list format from the LLM."""
    steps = []

    has_step_format = bool(re.search(r'Step\s+\d+\s*[:.]\s', procedure_text, re.IGNORECASE))

    if has_step_format:
        split_pat = r'(?=Step\s+\d+\s*[:.]\s)'
        header_pat = r'Step\s+(\d+)\s*[:.]\s*(.*?)(?:\n)'
    else:
        split_pat = r'(?:^|\n)(?=\d+\.\s)'
        header_pat = r'(\d+)\.\s*(.*?)(?:\n)'

    step_blocks = re.split(split_pat, procedure_text, flags=re.IGNORECASE)
    step_blocks = [b.strip() for b in step_blocks if b.strip()]

    print(f"DEBUG parse_steps: {len(step_blocks)} blocks (format: {'Step N' if has_step_format else 'N.'})")

    for idx, block in enumerate(step_blocks):
        header_match = re.match(header_pat, block, re.IGNORECASE)
        if not header_match:
            print(f"DEBUG: Block {idx} no header: {block[:80]}...")
            continue

        step_num = int(header_match.group(1))
        explanation = header_match.group(2).strip()
        explanation = re.sub(r'\*\*([^*]*)\*\*', r'\1', explanation)
        explanation = explanation.rstrip(':').strip()

        rest = block[header_match.end():]

        output_part = ''
        command_part = rest.strip()

        for pattern in [
            r'\n\s*\*{0,2}Output\*{0,2}\s*:\s*\n',
            r'\n\s*\*{0,2}Output\*{0,2}\s*:\s*',
            r'\*{0,2}Output\*{0,2}\s*:\s*\n',
        ]:
            parts = re.split(pattern, rest, maxsplit=1, flags=re.IGNORECASE)
            if len(parts) == 2:
                command_part = parts[0].strip()
                output_part = parts[1].strip()
                break
        else:
            print(f"DEBUG: Step {step_num} no Output: delimiter")

        command_part = re.sub(r'^\$\s*', '', command_part, flags=re.MULTILINE)
        command_part = re.sub(r'\*\*([^*]*)\*\*', r'\1', command_part)

        print(f"DEBUG: Step {step_num} cmd:{len(command_part)} out:{len(output_part)}")

        steps.append({
            'num': step_num,
            'explanation': explanation,
            'command': command_part,
            'output': output_part,
        })

    if not steps:
        print(f"DEBUG: No steps! Text: {procedure_text[:200]}...")
        steps = [{'num': 1, 'explanation': 'Execute the procedure', 'command': procedure_text, 'output': ''}]

    return steps


# ==================== API: Generate Content ====================
@app.route('/api/generate', methods=['POST'])
def api_generate():
    try:
        data = request.get_json()
        aim = data.get('aim', '')
        api_key = data.get('api_key', '')
        model = data.get('model', 'llama-3.3-70b-versatile')
        mode = data.get('mode', 'general')
        
        terminal_user = data.get('terminal_user', 'student')
        if not terminal_user.strip():
            terminal_user = 'student'
            
        terminal_host = data.get('terminal_host', 'kali')
        if not terminal_host.strip():
            terminal_host = 'kali'

        if not api_key:
            api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY not found in session or environment.")

        from groq import Groq
        client = Groq(api_key=api_key)

        if mode == 'os':
            prompt = f"""You are a professional Linux systems instructor preparing a practical lab file for a university Operating Systems course. Your instructor has assigned this aim:

"{aim}"

Work through the aim methodically, step by step, like an experienced professional would demonstrate in a real lab. Before each command or code block, write a clear explanation (1-2 lines) of what it does and why. Then write the command. Then show what the terminal actually displayed.

If the aim asks you to explore a command with its options, use it normally first, then show a few useful options — just like you'd actually try them in a real lab session. Don't robotically list every flag. Use your judgement.

If the aim asks for a C program (system calls, algorithms, etc.), write a clean, complete program. Compile and run it.

If the aim has multiple parts or multiple commands, work through each one properly.

Keep it natural. No filler. No padding. Just a real, useful practical.

Respond in this exact format:

[CONCEPT]
4-5 lines explaining the core OS concepts behind this practical in academic language. Keep it concise and focused. If there is a specific term or technique that needs extra explanation (e.g., what a system call is, what a process control block does), add ONE short follow-up paragraph for it — but only if truly needed. Do not over-explain obvious things.

[PROCEDURE]
Plain text, no markdown fences. Write it as numbered steps. Each step MUST include its own output immediately after the command.

Format each step EXACTLY like this:

Step 1: <what you're doing and why>
$ <command or code>
Output:
{terminal_user}@{terminal_host}:~$ <the command typed>
<realistic terminal output for this specific command>

Step 2: <what you're doing and why>
$ <command or code>
Output:
{terminal_user}@{terminal_host}:~$ <the command typed>
<realistic terminal output for this specific command>

...and so on. Only what the aim needs.

For C programs, write the full source code in one step (no $ prefix for the code itself), then compile and run as separate steps with $ prefix and their own outputs.

CRITICAL — Terminal output rules:
- The Output section for each step MUST start with the prompt and the command being typed on the FIRST line, then show the result below it. Do NOT add an extra prompt line at the end.
- Use this prompt: {terminal_user}@{terminal_host}:~$
- For root: root@{terminal_host}:~#
- Real permissions, real file sizes, real dates (Feb-Mar 2026), real PIDs, real kernel (6.1.0-18-amd64)
- No placeholders, no "...", no skipped output
- Each step's output must be realistic and complete

[CAPTION]
3-5 word caption for the experiment.
"""
        else:
            prompt = f"""You are an expert programming lab assistant. For this experiment aim:

"{aim}"

IMPORTANT GUIDELINES:
- Write clean and well-structured code.
- Provide a brief academic explanation of the core concepts being targeted.
- Show a realistic text output of running this code.
- If it's a programming language, provide the full source code.
- Do NOT include shell prompts like student@kali. Just show raw console execution outputs.

Respond EXACTLY in this format (use these exact tags):

[CONCEPT]
Write 3-4 lines explaining the concepts used. Academic style.

[CODE]
Write the code. Plain text only, no markdown fences.
Keep comments minimal — only where genuinely needed.

[OUTPUT]
Show REALISTIC output from running the code.
Make it look like a real terminal or console output. Do not show generic placeholder output.

[CAPTION]
Write a very short (3-5 words) descriptive caption for the output.
"""

        # Retry with backoff for rate limiting (429)
        max_retries = 3
        for attempt in range(max_retries + 1):
            try:
                chat_completion = client.chat.completions.create(
                    messages=[{'role': 'user', 'content': prompt}],
                    model=model,
                )
                text = chat_completion.choices[0].message.content
                break
            except Exception as api_err:
                err_str = str(api_err)
                if ('429' in err_str or 'rate' in err_str.lower()) and attempt < max_retries:
                    wait_time = 15 * (2 ** attempt)  # 15s, 30s, 60s
                    print(f"Rate limited. Waiting {wait_time}s before retry {attempt + 1}/{max_retries}...")
                    time.sleep(wait_time)
                else:
                    raise

        def extract_section(tag, text):
            pattern = rf"\[{tag}\](.*?)(?=\[(?:CONCEPT|CODE|PROCEDURE|OUTPUT|CAPTION)\]|$)"
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            
            if match:
                return match.group(1).strip()
            
            pattern_fallback = rf"(?:\*\*|##\s*)?{tag}(?:\*\*|:)?\s*\n(.*?)(?=\n(?:\*\*|##\s*)?(?:CONCEPT|CODE|PROCEDURE|OUTPUT|CAPTION)(?:\*\*|:)?\s*\n|\Z)"
            match_fallback = re.search(pattern_fallback, text, re.DOTALL | re.IGNORECASE)
            return match_fallback.group(1).strip() if match_fallback else None

        concept = extract_section("CONCEPT", text)
        caption = extract_section("CAPTION", text)
        if not concept: concept = "No concept description provided by API."
        if not caption: caption = "Experiment Output"

        result = {
            'concept': concept,
            'caption': caption,
            'mode': mode
        }

        if mode == 'os':
            procedure = extract_section("PROCEDURE", text) or "No procedure provided."
            procedure = re.sub(r'```[a-zA-Z]*', '', procedure).replace('```', '').strip()
            steps = parse_steps(procedure)
            result['steps'] = steps
            result['code'] = '\n\n'.join([f"Step {s['num']}: {s['explanation']}\n{s['command']}" for s in steps])
            result['output'] = '\n\n'.join([s['output'] for s in steps if s['output']])
        else:
            code = extract_section("CODE", text) or "// No code provided."
            output_part = extract_section("OUTPUT", text) or "No output provided."
            
            code = re.sub(r'```[a-zA-Z]*', '', code).replace('```', '').strip()
            output_part = re.sub(r'```', '', output_part).strip()

            result['code'] = code
            result['output'] = output_part

        return jsonify(result)
    except Exception as e:
        error_msg = str(e)
        status_code = 500
        if "401" in error_msg or "Invalid API Key" in error_msg or "Authentication" in error_msg:
            status_code = 401
        elif "429" in error_msg or "Rate limit" in error_msg:
            status_code = 429
        return jsonify({'error': error_msg}), status_code



# ==================== API: Download .docx ====================
def set_font(paragraph, font_name='Times New Roman', size=12, bold=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        run.bold = bold


def add_bold_para(doc, text, font_name='Times New Roman', size=12, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = True
    set_font(p, font_name=font_name, size=size, bold=True)
    return p


def add_labeled_para(doc, label, content, font_name='Times New Roman', size=12):
    p = doc.add_paragraph()
    if any(c in content for c in ['*', '•', '·']) or '  ' in content:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p.add_run(f'{label} ')
    run_label.bold = True
    run_label.font.name = font_name
    run_label.font.size = Pt(size)
    run_content = p.add_run(content)
    run_content.font.name = font_name
    run_content.font.size = Pt(size)
    return p


def add_code_para(doc, code_text, font_name='Times New Roman', size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.name = font_name
    run.font.size = Pt(size)


def add_caption_para(doc, text, experiment_no, step_no=None, font_name='Times New Roman', size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if step_no:
        label = f'Figure {experiment_no}.{step_no} - {text}'
    else:
        label = f'Figure {experiment_no} - {text}'
    run = p.add_run(label)
    run.font.name = font_name
    run.font.size = Pt(size)


def create_terminal_image(output_text, img_width=600):
    # Image settings
    width = img_width
    font_size = 16
    padding = 20
    
    # Check for font
    try:
        font = ImageFont.truetype("consola.ttf", font_size) # Windows console font
    except IOError:
        try:
            font = ImageFont.truetype("cour.ttf", font_size) # Courier
        except IOError:
            font = ImageFont.load_default()

    # Calculate height
    lines = str(output_text).split('\n')
    line_height = font_size + 9 
    height = (len(lines) * line_height) + (2 * padding)
    
    # Create Image
    img = Image.new('RGB', (width, height), color=(0, 0, 0)) # Pure black background
    d = ImageDraw.Draw(img)
    
    # Draw text
    y = padding
    for line in lines:
        try:
            text_line = line.replace('\r', '')
            # Normal font weight, matching Picture1.png color
            d.text((padding, y), text_line, font=font, fill=(201, 219, 213))
        except:
            pass
        y += line_height
        
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def add_normal_para(doc, text, font_name='Times New Roman', size=12, align=None):
    if align is None:
        if any(c in text for c in ['*', '•', '·']) or '  ' in text:
            align = WD_ALIGN_PARAGRAPH.LEFT
        else:
            align = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(p, font_name=font_name, size=size)
    return p


@app.route('/api/download', methods=['POST'])
def api_download():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data received for export.'}), 400
            
        experiments = data.get('experiments', [])
        if not experiments:
            return jsonify({'error': 'No experiment artifacts found to bundle.'}), 400
            
        settings = data.get('settings', {})

        font_name = settings.get('fontName', 'Times New Roman')
        body_size = int(settings.get('bodySize', 12))
        heading_size = int(settings.get('headingSize', 14))
        code_size = int(settings.get('codeSize', 10))
        caption_size = int(settings.get('captionSize', 10))
        image_width_inches = float(settings.get('imageWidth', 5.0))
        terminal_img_width = int(settings.get('terminalImgWidth', 600))
        output_filename = settings.get('outputFilename', 'Generated_Practical_File.docx')

        mode = data.get('mode', 'general')
        doc = Document()

        for i, exp in enumerate(experiments, 1):
            aim = exp.get('aim', 'N/A')
            concept = exp.get('concept', 'No concept description provided.')
            caption = exp.get('caption', 'Terminal Output Preview')
            steps = exp.get('steps', [])

            # Experiment heading
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'Experiment No. {i}')
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(heading_size)
            doc.add_paragraph('')

            add_labeled_para(doc, 'Aim:', aim, font_name, body_size)
            doc.add_paragraph('')
            
            if mode == 'os':
                add_bold_para(doc, 'Theory:', font_name, body_size)
                add_normal_para(doc, concept, font_name, body_size)
                doc.add_paragraph('')

                # Procedure with per-step output images
                add_bold_para(doc, 'Procedure:', font_name, body_size)

                if steps:
                    for step in steps:
                        step_num = step.get('num', '')
                        explanation = step.get('explanation', '')
                        command = step.get('command', '')
                        output = step.get('output', '')

                        # Step explanation
                        add_normal_para(doc, f"Step {step_num}: {explanation}", font_name, body_size)
                        
                        # For multi-line code (C programs etc.), show source code as text
                        is_multiline_code = command.count('\n') > 2
                        if is_multiline_code:
                            add_code_para(doc, command, font_name, code_size)

                        # Output as terminal image
                        if output.strip():
                            try:
                                img_buf = create_terminal_image(output, terminal_img_width)
                                pic_para = doc.add_paragraph()
                                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = pic_para.add_run()
                                run.add_picture(img_buf, width=Inches(image_width_inches))
                                add_caption_para(doc, explanation, i, step_num, font_name, caption_size)
                            except Exception as img_err:
                                print(f"DEBUG: Step {step_num} image error: {img_err}")
                                add_code_para(doc, output, font_name, code_size)
                        
                        doc.add_paragraph('')  # spacing between steps
                else:
                    # Fallback
                    code = exp.get('code', '// No procedure available.')
                    output = exp.get('output', 'No output.')
                    add_code_para(doc, code, font_name, code_size)
                    doc.add_paragraph('')
                    add_bold_para(doc, 'Output:', font_name, body_size)
                    try:
                        img_buf = create_terminal_image(output, terminal_img_width)
                        pic_para = doc.add_paragraph()
                        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = pic_para.add_run()
                        run.add_picture(img_buf, width=Inches(image_width_inches))
                        add_caption_para(doc, caption, i, None, font_name, caption_size)
                    except Exception as img_err:
                        print(f"DEBUG: Error creating terminal image: {img_err}")
                        add_code_para(doc, output, font_name, code_size)

            else:
                code = exp.get('code', '// No code available.')
                output = exp.get('output', 'Program executed successfully.')

                add_labeled_para(doc, 'Concept Used:', concept, font_name, body_size)
                doc.add_paragraph('')
                add_bold_para(doc, 'Code:', font_name, body_size)
                add_code_para(doc, code, font_name, code_size)
                doc.add_paragraph('')
                add_bold_para(doc, 'Output:', font_name, body_size)

                try:
                    img_buf = create_terminal_image(output, terminal_img_width)
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = pic_para.add_run()
                    run.add_picture(img_buf, width=Inches(image_width_inches))
                    add_caption_para(doc, caption, i, None, font_name, caption_size)
                except Exception as img_err:
                    print(f"DEBUG: Error creating terminal image: {img_err}")
                    add_normal_para(doc, f'[Visual Output Unavailable - Log Trace follows]', font_name, body_size)
                    add_code_para(doc, output, font_name, code_size)

            if i < len(experiments):
                doc.add_page_break()

        file_buf = io.BytesIO()
        doc.save(file_buf)
        file_buf.seek(0)

        return send_file(
            file_buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_filename
        )
    except Exception as e:
        print(f"CRITICAL EXPORT ERROR: {traceback.format_exc()}")
        return jsonify({'error': f'Export Pipeline Fault: {str(e)}'}), 500


if __name__ == '__main__':
    print('\n  ⚡ PractiGen running at http://localhost:5000\n')
    app.run(host='0.0.0.0', port=5000, debug=True)
